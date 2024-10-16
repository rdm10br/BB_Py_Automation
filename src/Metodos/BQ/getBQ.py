import docx.opc
import docx.opc.exceptions
import regex as re
import docx, spacy
from spacy.matcher import Matcher
from functools import lru_cache
from docx.shared import RGBColor

nlp = spacy.load("pt_core_news_sm")
matcher = Matcher(nlp.vocab)

@lru_cache
def read_document(path: str) -> str:
    '''
    Return the file content
    
    Function to read a docx file, given the file path in the ```path```
    variable, and store in a variable
    '''
    try:
        doc = docx.Document(docx=path)
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        text_content = "\n".join(content)
        text_content = re.sub(r'\n+', '', text_content)
        if text_content == '':
            try:
                tables_content = []
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text not in row_data:
                                row_data.append(cell_text)
                        if row_data:
                            table_data.append("\t".join(row_data))
                    if table_data:
                        tables_content.append("\n".join(table_data))
                    text_content = "\n".join(tables_content)
                return text_content
            except Exception as e:
                print(f"Error reading document: {e}")
                return None
        return "\n".join(content)
  
    except Exception as e:
        print(f"Error reading document: {e}")
        return None

@lru_cache
def enunciado_count (path: str) -> int:
    """
    Return how many statments on the file

    Args:
        path (str): Doc Path to the Questionary

    Returns:
        int: mathes - how many Statements this questionary have
    """
    texto = read_document(path)
    if texto is None:
        return 0
    if texto == '':
        return 0
    
    doc = nlp(texto)
    
    pattern = [{"LEMMA": "Questão"}, {"IS_DIGIT": True}]
    matcher.add("Questions", [pattern])
    
    matches = len(matcher(doc))
    
    if matches == 0:
        try:
            pattern = [{"LEMMA": "QUESTÃO"}, {"IS_DIGIT": True}]
            matcher.add("QUESTIONS", [pattern])
            matches = len(matcher(doc))
            return matches
        except:
            return matches
    
    return matches

def extract_text_between_markers(text: str, start_marker: str, end_marker: str):
    """
    Function to extract text between two markers using regular expressions.
    """
    try:
        start_index = text.find(start_marker)
        # print(start_index+1)
        if start_index == -1:
            return ''
    except:
        return 'index error, question not found'

    # Only consider the text after the start_marker
    text_after_start = text[start_index + len(start_marker):]
    # text_after_start.
    # pattern = re.compile(rf'(.*?){re.escape(end_marker)}', re.DOTALL)
    pattern = re.compile(rf'(.*?)(?=^\s*{re.escape(end_marker)})', re.DOTALL | re.MULTILINE)
    match = re.search(pattern, text_after_start)
    if match:
        extracted_text = match.group(1).strip()
        
        # cleaned_text = re.sub(r'\s+', ' ', extracted_text)

        # # Remove excessive new lines
        # cleaned_text = re.sub(r'\n+', '\n', cleaned_text)

        # # Alternatively, you can combine both into a single operation:
        # cleaned_text = re.sub(r'\s+', ' ', extracted_text) # Step 1: Replace multiple spaces with a single space
        # cleaned_text = re.sub(r'\n+', '\n', cleaned_text)
        return extracted_text
    else:
        return ""

@lru_cache
def get_enunciados(filename: str):
    
    text = read_document(filename)
    if text is None:
        return []
    
    q = enunciado_count(filename)
    question = []
    
    for i in range(q):
        i += 1
        start_marker = f'Questão {i}'
        end_marker = r'a)'
        extracted_text = extract_text_between_markers(text, start_marker, end_marker)
        if extracted_text == '':
            start_marker = f'QUESTÃO {i}'
            end_marker = r'a)'
            extracted_text = extract_text_between_markers(text, start_marker, end_marker)
        question.append(extracted_text)
    return question

def get_Enunciado(index: int, path: str) -> str:
    '''
    Return question statement
    
    Function to get the statement from the ```path``` file you get
    the ```index```+1 question
    '''
    # doc = read_document(path=path)
    # enunciado = re.findall(pattern=regex_Enunciado, string=doc).copy()[index]
    # return enunciado
    question = get_enunciados(filename=path)
    try:
        result = question[index]
        return result
    except:
        print('Index out of Range or Question not found!')
        return ''

@lru_cache
def get_Alternativa(index: int, path: str, choices: str) -> str:
    '''
    Return question choices
    
    Function to get the choices from the ```path``` file you get
    the ```index```+1 question and ```choices``` given in the method
    '''
    doc = read_document(path=path)
    index_marker = doc.find(get_Enunciado(index=index, path=path))
    doc = doc[index_marker:]
    if doc is None:
        return ""
    match choices.upper():
        case 'A':
            start_marker = get_Enunciado(index=index, path=path)
            end_marker = r'b)'
            match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
            if match:
                cleaned_text = re.sub(r'^[a-e]\)\s*', '', match, flags=re.MULTILINE)
                return cleaned_text
        case 'B':
            start_marker = get_Alternativa_hole(index=index, path=path, choices='a')
            end_marker = r'c)'
            match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
            if match:
                cleaned_text = re.sub(r'^[a-e]\)\s*', '', match, flags=re.MULTILINE)
                return cleaned_text
        case 'C':
            start_marker = get_Alternativa_hole(index=index, path=path, choices='b')
            end_marker = r'd)'
            match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
            if match:
                cleaned_text = re.sub(r'^[a-e]\)\s*', '', match, flags=re.MULTILINE)
                return cleaned_text
        case 'D':
            start_marker = get_Alternativa_hole(index=index, path=path, choices='c')
            end_marker = r'e)'
            match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
            if match:
                cleaned_text = re.sub(r'^[a-e]\)\s*', '', match, flags=re.MULTILINE)
                return cleaned_text
        case 'E':
            start_marker = get_Alternativa_hole(index=index, path=path, choices='d')
            # end_marker = r'Justificativa|Gabarito|Questão|\d\.|\d+\.'
            end_marker = r'Justificativa'
            match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
            if match:
                cleaned_text = re.sub(r'^[a-e]\)\s*', '', match, flags=re.MULTILINE)
                return cleaned_text
            elif match == '':
                end_marker = r'ALTERNATIVA CORRETA'
                match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
                if match:
                    cleaned_text = re.sub(r'^[a-e]\)\s*', '', match, flags=re.MULTILINE)
                    return cleaned_text
                elif match == '':
                    end_marker = r'JUSTIFICATIVA'
                    match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
                    if match:
                        cleaned_text = re.sub(r'^[a-e]\)\s*', '', match, flags=re.MULTILINE)
                        return cleaned_text
                    elif match == '':
                        ...
        case _ :
            print('''Por favor verifique a chamada da função get_Alternativa,
                  tipo de alternativa desejada não esperada pela função''')
            
@lru_cache
def get_Alternativa_hole(index: int, path: str, choices: str) -> str:
    '''
    Return question choices
    
    Function to get the choices from the ```path``` file you get
    the ```index```+1 question and ```choices``` given in the method
    '''
    doc = read_document(path=path)
    index_marker = doc.find(get_Enunciado(index=index, path=path))
    doc = doc[index_marker:]
    if doc is None:
        return ""
    match choices.upper():
        case 'A':
            start_marker = get_Enunciado(index=index, path=path)
            end_marker = r'b)'
            match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
            if match:
                return match
        case 'B':
            start_marker = get_Alternativa_hole(index=index, path=path, choices='a')
            end_marker = r'c)'
            match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
            if match:
                return match
        case 'C':
            start_marker = get_Alternativa_hole(index=index, path=path, choices='b')
            end_marker = r'd)'
            match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
            if match:
                return match
        case 'D':
            start_marker = get_Alternativa_hole(index=index, path=path, choices='c')
            end_marker = r'e)'
            match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
            if match:
                return match
        case 'E':
            start_marker = get_Alternativa_hole(index=index, path=path, choices='d')
            # end_marker = r'Justificativa|Gabarito|Questão|\d\.|\d+\.'
            end_marker = r'Justificativa'
            match = extract_text_between_markers(text=doc, start_marker=start_marker, end_marker=end_marker)
            if match:
                return match
        case _ :
            print('''Por favor verifique a chamada da função get_Alternativa,
                  tipo de alternativa desejada não esperada pela função''')

@lru_cache
def get_correct_alternative_by_any_color(path: str) -> str:
    try:
        document = docx.Document(path)
        alternativas = []
        alternativas_com_destaque = []
        alternativas_corretas = []
    
        for paragraph in document.paragraphs:
            texto = paragraph.text.strip()
            if texto and texto[0] in "ABCDEabcde" and texto[1] in ").":
                alternativas.append(texto)
                for run in paragraph.runs:
                    if run.font.highlight_color or (run.font.color and (run.font.color.rgb == RGBColor(255, 0, 0))):
                        alternativas_com_destaque.append(texto)
                        break
        
        if len(alternativas_com_destaque) == 0:
            for paragraph in document.paragraphs:
                texto = paragraph.text.strip()
                # Gabarito: A
                if texto and texto in "Gabarito: ":
                    alternativas_corretas.append(texto[-1])
                # Resposta: b
                elif texto and texto in "Resposta: ":
                    alternativas_corretas.append(texto[-1])
                # Alternativa Correta: letra A.
                elif texto and texto in "Alternativa Correta: letra ":
                    alternativas_corretas.append(texto[-2])
            
            # ALTERNATIVA CORRETA 12%
            
            alternativas_com_destaque = alternativas_corretas
        
        return alternativas_com_destaque

    except (docx.opc.exceptions.PackageNotFoundError, IndexError) as e:
        print(f"\nError: {e}")
        return None

    except Exception as e:
        print(f"\nUnexpected error: {e}")
        return None
    
@lru_cache
def get_correct_alternative_from_list (path: str, index: int):
    awnser_list = get_correct_alternative_by_any_color(path)
    return awnser_list[index][0]

def main() -> None:
    path = r'C:\Users\013190873\Downloads\Questionário_Álgebra Linear_Unidade I_DIGITAL PAGES_ORIGINAL.docx'
    
    # teste = enunciado_count(path=path)
    # print(f'\n Enunciado count: {teste}')
    
    index = 19
    teste2 = get_Enunciado(index=index, path=path)
    
    print(f'\n Question:\n{teste2}')
    
    teste3 = get_Alternativa(index=index, path=path, choices='a')
    
    print(f'\n Choices:\n{teste3}')
    
    teste_right = get_correct_alternative_from_list(path=path, index=index)
    print(f'\nalternativa correta: {teste_right}\n')

if __name__ == "__main__":
    main()